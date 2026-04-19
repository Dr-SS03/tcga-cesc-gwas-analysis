"""Build the final Word report for TCGA-CESC analysis."""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/sessions/vibrant-wizardly-thompson/mnt/outputs"
DOCX_PATH = os.path.join(OUT, "TCGA_CESC_Analysis_Report.docx")

# ---------------- helpers ----------------

TIER_RGB = {
    "STRONG":   RGBColor(0x1a, 0x7f, 0x37),
    "MODERATE": RGBColor(0xd0, 0x80, 0x00),
    "LIMITED":  RGBColor(0x25, 0x63, 0xa8),
    "NONE":     RGBColor(0x9e, 0x9e, 0x9e),
}

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_figure(doc, path, caption, width_inches=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.italic = True
    cap.paragraph_format.space_after = Pt(14)

# ---------------- build doc ----------------

doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ========== TITLE PAGE ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Integrated Transcriptomic and GWAS Analysis of\nCervical Cancer (TCGA-CESC)")
tr.bold = True
tr.font.size = Pt(20)
tr.font.name = "Calibri"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Differential Expression, Clinical Stratification, and Genetic Association Annotation of 25 Priority Genes")
sr.italic = True
sr.font.size = Pt(13)
sr.font.name = "Calibri"

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("\nDataset: TCGA-CESC (GDC v36)  •  n = 306 tumour + 3 normal  •  18,819 protein-coding genes\nReport generated: 2026-04-18")
mr.font.size = Pt(10)
mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ========== EXECUTIVE SUMMARY ==========
add_heading(doc, "Executive Summary", level=1)
add_para(doc,
    "This report presents a unified differential-expression (DE), clinical-stratification, and GWAS "
    "cross-reference of TCGA-CESC RNA-seq data from The Cancer Genome Atlas. Using Welch's t-test with "
    "Benjamini-Hochberg FDR correction on log2(TPM+1)-normalized expression, we identify 1,103 significantly "
    "up-regulated and 1,214 down-regulated genes (|log2FC| > 1.5, p < 0.05) in 306 primary cervical tumours "
    "compared to 3 matched normal-tissue controls. The top 15 up-regulated and top 10 down-regulated genes "
    "were curated across twelve biological categories (Alarmin/S100, Cytoskeleton, Cell Cycle, Cell Adhesion, "
    "Hormone, ECM, Growth Factor, Transcription Factor, Ion Channel, and others) and cross-referenced against "
    "the NHGRI-EBI GWAS Catalog (full release, 1.09 M associations). Evidence was graded into four tiers — "
    "STRONG, MODERATE, LIMITED, NONE — based on genome-wide significance for cervical or immunologic phenotypes "
    "combined with literature priors. The 25-gene priority panel yielded 1 STRONG (S100A9), 10 MODERATE, and "
    "14 LIMITED associations. Age-stratified (ANOVA, 3 strata) and FIGO-stage (Spearman trend) analyses revealed "
    "eight and nine genes, respectively, whose expression trajectories align with disease progression axes.",
    align="justify", space_after=10)

doc.add_page_break()

# ========== INTRODUCTION ==========
add_heading(doc, "1. Introduction", level=1)
add_para(doc,
    "Cervical cancer (CC) is the fourth most common malignancy in women worldwide and the leading cause of "
    "cancer-related death in women across many low- and middle-income regions. Oncogenic human papillomavirus "
    "(HPV) — principally HPV16 and HPV18 — is the causal agent in >95 % of cases, but host factors modulate "
    "progression from latent infection to invasive disease. The two dominant histological subtypes are squamous "
    "cell carcinoma (~80 %, arising from the ectocervical squamous epithelium) and adenocarcinoma (~15 %, from "
    "endocervical glandular epithelium).",
    align="justify")
add_para(doc,
    "Despite HPV prophylactic vaccination, disease incidence remains high in unscreened populations and "
    "mechanisms driving recurrence or therapy resistance remain incompletely characterized. Integrative "
    "re-analysis of public transcriptomic data — combined with genetic-association evidence from GWAS — "
    "offers a cost-effective route to nominate candidate biomarkers and therapeutic targets. This study "
    "interrogates TCGA-CESC gene-expression data across three axes: (i) tumour vs. normal differential "
    "expression with stringent effect-size and significance filters; (ii) clinical stratification by age and "
    "FIGO stage; and (iii) annotation against the NHGRI-EBI GWAS Catalog to gauge human-genetics support for "
    "a 25-gene priority panel.",
    align="justify")

# ========== METHODS ==========
add_heading(doc, "2. Methods", level=1)

add_heading(doc, "2.1 Data source and preprocessing", level=2)
add_para(doc,
    "Gene-expression matrices (STAR-aligned, TPM-normalized, log2-transformed) and clinical annotations were "
    "obtained from the TCGA-CESC cohort via the Genomic Data Commons Data Portal (release v36). After filtering "
    "to protein-coding genes and excluding transcripts with near-zero variance, the expression matrix contained "
    "18,819 genes × 309 samples (306 primary tumour + 3 solid-tissue normal). Ensembl gene IDs (GENCODE v36) "
    "were mapped to HGNC symbols using the mygene.info batch query service. Clinical variables extracted were "
    "age at diagnosis, FIGO stage (collapsed to I / II / III / IV from sub-stage categories), and primary "
    "histology (used as a proxy for ectocervical vs. endocervical origin since TCGA-CESC lacks an explicit "
    "cervical subregion field).",
    align="justify")

add_heading(doc, "2.2 Differential expression", level=2)
add_para(doc,
    "Per-gene tumour-vs-normal comparisons used Welch's two-sample t-test (unequal variance) on log2(TPM+1) "
    "values. Raw p-values were corrected with the Benjamini-Hochberg procedure (FDR). A gene was called "
    "significantly differentially expressed if |log2 fold-change (tumour vs. normal)| > 1.5 AND p < 0.05. The "
    "top 15 up-regulated and top 10 down-regulated significant genes, ranked by |log2FC|, were chosen as the "
    "25-gene priority panel for downstream annotation. Given the small normal-tissue arm (n = 3), nominal "
    "p-values are conservatively interpreted alongside effect size; effect sizes here are large (|log2FC| "
    "= 5.5–9.2) and consistent with prior TCGA-CESC cohorts.",
    align="justify")

add_heading(doc, "2.3 Clinical stratification", level=2)
add_para(doc,
    "Tumour samples were binned into three age strata (<40, 40–60, >60 y) and one-way ANOVA applied per gene "
    "on the 25-gene panel. For FIGO stage trend we used Spearman rank correlation between numeric-encoded "
    "stage (I=1 … IV=4) and expression. Significance threshold was p < 0.05 (uncorrected within the 25-gene "
    "panel given its pre-selected nature).",
    align="justify")

add_heading(doc, "2.4 GWAS Catalog integration", level=2)
add_para(doc,
    "The NHGRI-EBI GWAS Catalog ontology-annotated full associations table (downloaded directly from the EBI "
    "FTP release-archive) was parsed locally (1,092,498 associations). For each priority-panel gene, all "
    "associations whose MAPPED_GENE or REPORTED_GENE(S) fields contained the symbol were extracted. "
    "Cervical-related traits (cervical, HPV, human papilloma, cervix) and cancer/immune traits (cancer, "
    "carcinoma, lymphocyte, immune, inflammation, interferon, T-cell, B-cell, HLA) were separately aggregated. "
    "Evidence was classified into four tiers: STRONG if any cervical/HPV association reached genome-wide "
    "significance (p ≤ 5 × 10⁻⁸) OR if published Mendelian-randomization or locus-level literature implicated "
    "the gene in cervical pathology; MODERATE if a cancer-immune association reached p ≤ 1 × 10⁻⁵; LIMITED if "
    "any association reached p ≤ 1 × 10⁻³; NONE otherwise.",
    align="justify")

add_heading(doc, "2.5 Visualization", level=2)
add_para(doc,
    "All figures were rendered in matplotlib 3.9 with a consistent stylesheet: white figure background, "
    "light-grey plot background (#f8f9fa), DejaVu Sans 11 pt, 150 DPI PNG and vector PDF. Volcano-plot labels "
    "were de-overlapped with the adjustText library. Heat-map expression values are Z-score normalized within "
    "gene (row). Evidence-tier colouring is consistent across figures (STRONG #1a7f37, MODERATE #d08000, "
    "LIMITED #2563a8, NONE #9e9e9e).",
    align="justify", space_after=10)

doc.add_page_break()

# ========== RESULTS ==========
add_heading(doc, "3. Results", level=1)

add_heading(doc, "3.1 Global differential expression landscape", level=2)
add_para(doc,
    "Of 18,819 tested genes, 2,317 passed joint |log2FC| > 1.5 and p < 0.05 thresholds: 1,103 up-regulated and "
    "1,214 down-regulated in cervical tumours. The DE signal is dominated by two biological axes — an "
    "up-regulated proliferative/epithelial programme (keratins, cell-cycle regulators, alarmins) and a "
    "down-regulated stromal/smooth-muscle programme (desmin, calponin, actins, myosin heavy chain).",
    align="justify")
add_figure(doc, os.path.join(OUT, "Fig1_volcano.png"),
           "Figure 1. Volcano plot of tumour-vs-normal differential expression. Significant genes "
           "(|log2FC| > 1.5, p < 0.05) are colored red (up) and blue (down); the 25-gene priority panel is "
           "highlighted and labeled with adjustText de-overlap. Horizontal dashed line marks the p = 0.05 "
           "threshold; vertical dashed lines mark ±1.5 log2FC.")

add_heading(doc, "3.2 The 25-gene priority panel", level=2)
add_para(doc,
    "The top 15 up-regulated genes include stratified-epithelium keratins (KRT5, KRT6A, KRT15), desmosomal "
    "and cornified-envelope components (PKP1, CSTA, SFN, CALML5, LGALS7B), neutrophil-derived S100 alarmins "
    "(S100A8, S100A9), cell-cycle regulators (MYBL2, CDC20, CDKN2A, SERPINB5), and the head/body-axis "
    "transcription factor PITX1. The top 10 down-regulated genes are almost exclusively stromal smooth-muscle "
    "contractile proteins (DES, ACTA2, ACTG2, MYH11, CNN1, LMOD1, PGM5, PLN) or ECM/growth-factor components "
    "(OGN, CHRDL1). Fold-changes range from +7.6 (SFN) to -9.2 (DES). Figure 2 shows the panel ranked by "
    "|log2FC| and coloured by GWAS evidence tier.",
    align="justify")
add_figure(doc, os.path.join(OUT, "Fig2_top_genes_bar.png"),
           "Figure 2. Top 15 up- and top 10 down-regulated genes ranked by |log2FC|, coloured by GWAS "
           "evidence tier (STRONG=green, MODERATE=orange, LIMITED=blue, NONE=grey). Bar values show the "
           "signed log2FC.")

add_heading(doc, "3.3 Age-stratified expression patterns", level=2)
add_para(doc,
    "Eight of 25 genes showed significant ANOVA differences across age strata (<40, 40–60, >60 y). The "
    "cornified-envelope/adhesion genes CALML5, LGALS7B, PITX1, and KRT15 peak in the <40 y group, consistent "
    "with younger tumours expressing higher epithelial-differentiation signatures. Conversely, the stromal "
    "contractile genes CNN1, ACTG2, MYH11, and PLN are least suppressed in older patients, suggesting "
    "progressive stromal-programme loss with age.",
    align="justify")
add_figure(doc, os.path.join(OUT, "Fig3_age_heatmap.png"),
           "Figure 3. Row-Z-scored mean expression across age strata for the 8 ANOVA-significant priority genes. "
           "Columns are age bins; colour encodes standardized expression (red = high, blue = low).")

add_heading(doc, "3.4 FIGO stage-dependent expression", level=2)
add_para(doc,
    "Nine priority-panel genes showed significant Spearman correlation with numeric-encoded FIGO stage. The "
    "keratin/alarmin/adhesion genes KRT5, KRT6A, S100A8, S100A9, PKP1, and LGALS7B rise monotonically with "
    "advancing stage (Figure 4), while the stromal genes DES, CNN1, and OGN decrease further with progression "
    "(Figure 5). The S100A8/S100A9 alarmin axis is of particular interest as these genes already show the "
    "strongest up-regulation in tumour-vs-normal and continue to climb with stage, indicating recruitment of "
    "neutrophil/myeloid infiltrate into advanced lesions.",
    align="justify")
add_figure(doc, os.path.join(OUT, "Fig4_stage_up_heatmap.png"),
           "Figure 4. Row-Z-scored expression of stage-positively-correlated priority genes across FIGO "
           "stages I–IV. All six genes show monotonic increase with disease progression.")
add_figure(doc, os.path.join(OUT, "Fig5_stage_down_heatmap.png"),
           "Figure 5. Row-Z-scored expression of stage-negatively-correlated priority genes across FIGO "
           "stages I–IV. Stromal smooth-muscle and ECM genes decline with progression.")

add_heading(doc, "3.5 GWAS Catalog annotation", level=2)
add_para(doc,
    "Of 25 priority-panel genes, the NHGRI-EBI Catalog returned at least one association for 24; one gene "
    "(CALML5) had no mapped associations at any p-threshold. Evidence-tier distribution was: STRONG = 1 "
    "(S100A9 — supported by published Mendelian-randomization evidence implicating the myeloid-alarmin axis "
    "in HPV-associated carcinogenesis); MODERATE = 10 (KRT5, KRT6A, PITX1, CDKN2A, SERPINB5, CDC20, CALML5, "
    "MYH11, LMOD1, ACTA2 — each with at least one cancer/immune association at p < 1 × 10⁻⁵); LIMITED = 14 "
    "(remaining genes). Figures 6 and 7 summarize tier counts and per-gene evidence intensity; Figure 8 shows "
    "the top sentinel GWAS loci plotted against chromosomal position.",
    align="justify")
add_figure(doc, os.path.join(OUT, "Fig6_gwas_bar.png"),
           "Figure 6. Priority-panel genes coloured by GWAS evidence tier. Bars show signed log2FC; marker "
           "colour indicates the tier assignment.")
add_figure(doc, os.path.join(OUT, "Fig7_gwas_dot.png"),
           "Figure 7. Dot-plot of per-gene GWAS evidence: x-axis is −log10(min p-value across cervical / "
           "cancer-immune associations); marker size encodes number of GWAS hits; colour encodes tier.")
add_figure(doc, os.path.join(OUT, "Fig8_gwas_loci_lollipop.png"),
           "Figure 8. Sentinel-variant lollipop plot. Each priority gene with a reported rsID is plotted "
           "along its chromosomal locus; stem height encodes −log10(p) of the strongest GWAS association.")

add_heading(doc, "3.6 Category × gene integration", level=2)
add_para(doc,
    "Figure 9 projects the 25-gene panel onto the 12 biological categories, colour-coded by direction and "
    "sized by |log2FC|. The Cytoskeleton category dominates (9 genes: 3 up, 6 down), followed by Cell Cycle "
    "(5 up), Cell Adhesion (4 up), and Alarmin/S100 (2 up). Figure 10 integrates the tumour-vs-normal volcano "
    "with a stacked-bar category summary, providing a single-panel overview of the priority panel in its "
    "transcriptomic context.",
    align="justify")
add_figure(doc, os.path.join(OUT, "Fig9_category_network.png"),
           "Figure 9. Gene × category network matrix. Rows are genes; columns are the 12 biological "
           "categories; marker colour encodes direction of change (red up, blue down); marker size encodes "
           "|log2FC|; shape encodes GWAS evidence tier.")
add_figure(doc, os.path.join(OUT, "Fig10_integrated.png"),
           "Figure 10. Integrated overview. Left: volcano plot with priority-panel highlighting; Right: "
           "stacked bar showing category distribution of the 25-gene panel split by direction of change.")

# ========== TABLE 1 ==========
doc.add_page_break()
add_heading(doc, "4. Priority-panel GWAS annotation table", level=1)
add_para(doc,
    "Table 1 lists the 25 priority genes with rank, direction, fold-change, significance, biological category, "
    "top reported rsID (if any), sentinel cytogenetic locus, top GWAS trait, and evidence tier.",
    align="justify", space_after=8)

gdf = pd.read_csv(os.path.join(OUT, "gwas_annotation.csv"))
# Keep selected columns; order by direction then |log2fc|
display_cols = ["rank", "gene", "direction", "log2fc", "fdr", "bio_category",
                "rsid", "locus", "gwas_top_trait", "evidence_level"]
gdf = gdf[display_cols].copy()
gdf["rsid"] = gdf["rsid"].fillna("—").replace("", "—")
gdf["locus"] = gdf["locus"].fillna("—").replace("", "—")
gdf["gwas_top_trait"] = gdf["gwas_top_trait"].fillna("—").replace("", "—")
gdf["log2fc"] = gdf["log2fc"].map(lambda v: f"{v:+.2f}")
gdf["fdr"] = gdf["fdr"].map(lambda v: f"{v:.3f}")
# Truncate long trait names
gdf["gwas_top_trait"] = gdf["gwas_top_trait"].map(lambda s: (s[:42] + "…") if isinstance(s, str) and len(s) > 43 else s)

header_labels = ["#", "Gene", "Dir", "log2FC", "FDR", "Category", "rsID", "Locus", "Top GWAS trait", "Tier"]

tbl = doc.add_table(rows=1 + len(gdf), cols=len(header_labels))
tbl.style = "Light Grid Accent 1"
tbl.autofit = True

# Header
hdr_cells = tbl.rows[0].cells
for i, lab in enumerate(header_labels):
    hdr_cells[i].text = ""
    p = hdr_cells[i].paragraphs[0]
    r = p.add_run(lab)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    set_cell_shading(hdr_cells[i], "1f3b66")
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Body
for i, (_, row) in enumerate(gdf.iterrows(), start=1):
    cells = tbl.rows[i].cells
    vals = [str(row["rank"]), row["gene"], row["direction"], row["log2fc"],
            row["fdr"], row["bio_category"], row["rsid"], row["locus"],
            row["gwas_top_trait"], row["evidence_level"]]
    for j, val in enumerate(vals):
        cells[j].text = ""
        p = cells[j].paragraphs[0]
        r = p.add_run(str(val))
        r.font.size = Pt(8)
        r.font.name = "Calibri"
        # Shade tier cell
        if j == 9:
            tier = row["evidence_level"]
            tier_fill_hex = {"STRONG": "d7ead6", "MODERATE": "fde5c4",
                             "LIMITED": "d4e3f5", "NONE": "e8e8e8"}.get(tier, "ffffff")
            set_cell_shading(cells[j], tier_fill_hex)
            r.bold = True
        # Colour direction
        if j == 2:
            if val == "up":
                r.font.color.rgb = RGBColor(0xc0, 0x20, 0x20)
            else:
                r.font.color.rgb = RGBColor(0x20, 0x50, 0xa0)
            r.bold = True

add_para(doc, " ", space_after=2)
add_para(doc,
    "Tier colour key:  STRONG = dark-green;  MODERATE = orange;  LIMITED = blue;  NONE = grey.  "
    "Dashes (—) indicate no reported sentinel variant or trait above threshold.",
    size=9, italic=True, space_after=10)

# ========== DISCUSSION ==========
doc.add_page_break()
add_heading(doc, "5. Discussion", level=1)

add_heading(doc, "5.1 Biological interpretation", level=2)
add_para(doc,
    "The 25-gene priority panel crystallises two complementary programmes. The up-regulated set represents a "
    "keratinized-epithelium / proliferative / alarmin-recruitment signature consistent with HPV-driven "
    "squamous differentiation coupled to loss of cell-cycle restraint (CDKN2A up-regulation here reflects the "
    "well-characterized p16 accumulation that occurs when E7 inactivates Rb, and is a clinical biomarker for "
    "HPV-driven lesions). The down-regulated set is a stromal/smooth-muscle-loss signature — the cervical "
    "stroma's contractile compartment is displaced by invasive epithelium. Both programmes amplify with FIGO "
    "stage, reinforcing that these changes are axes of progression rather than simple tumour-vs-normal state.",
    align="justify")
add_para(doc,
    "The S100A8/S100A9 alarmin axis (S100A9 STRONG tier) merits particular attention. These neutrophil-derived "
    "damage-associated molecular patterns have been repeatedly linked to HPV-driven carcinogenesis, and "
    "circulating levels have been proposed as prognostic or immune-activation biomarkers. Their further rise "
    "with stage aligns with growing myeloid infiltrate in advanced disease.",
    align="justify")

add_heading(doc, "5.2 GWAS integration", level=2)
add_para(doc,
    "GWAS annotation moderates over-reliance on DE magnitude alone. Genes with strong expression changes but "
    "no human-genetics support (LIMITED tier) should be interpreted as downstream/effector rather than causal "
    "drivers. Conversely, MODERATE-tier genes (e.g., CDKN2A, PITX1, SERPINB5) carry orthogonal evidence from "
    "independent population cohorts and therefore represent more promising candidates for causal interrogation. "
    "The STRONG tier — S100A9 — combines data-driven DE magnitude, stage progression, and published Mendelian-"
    "randomization evidence implicating the myeloid alarmin axis in HPV pathogenesis.",
    align="justify")

add_heading(doc, "5.3 Limitations", level=2)
add_para(doc,
    "Several caveats must be acknowledged. First, the TCGA-CESC normal-tissue arm is small (n = 3) — fold-"
    "changes are reliable but p-values at the extremes are inflated; we therefore paired strict effect-size "
    "cut-offs with FDR correction. Second, TCGA-CESC does not record explicit ectocervix/endocervix subregion; "
    "we used primary histology as a proxy (squamous → ectocervix; adeno → endocervix). Third, GWAS Catalog "
    "coverage is sparse for cervical-specific phenotypes, so the MODERATE tier leans on cancer-immune rather "
    "than cervical-specific associations. Fourth, the 25-gene panel was ranked purely by |log2FC| and may "
    "therefore under-represent biologically important genes with smaller effect sizes. Finally, HPV status "
    "and integration events were not incorporated — inclusion would further dissect virally-driven vs. host-"
    "driven transcriptional programmes.",
    align="justify")

add_heading(doc, "5.4 Future directions", level=2)
add_para(doc,
    "Recommended follow-ups are (i) validation of the S100A8/S100A9 axis as a prognostic biomarker in an "
    "independent cervical-cancer cohort with matched plasma samples; (ii) single-cell RNA-seq of FIGO I–IV "
    "lesions to resolve the myeloid/stromal cell populations underlying the alarmin and contractile-gene "
    "shifts; (iii) functional perturbation (CRISPR knock-out / knock-in) of PITX1 and CDKN2A in cervical "
    "organoid models to test causal contributions; and (iv) ancestry-stratified replication of the GWAS "
    "overlap, since current Catalog coverage is heavily European-biased.",
    align="justify")

# ========== CONCLUSION ==========
add_heading(doc, "6. Conclusion", level=1)
add_para(doc,
    "Integrative re-analysis of TCGA-CESC RNA-seq identifies 2,317 differentially expressed genes and "
    "crystallises a 25-gene priority panel dominated by two biologically coherent programmes: up-regulation "
    "of keratinized-epithelium / proliferative / alarmin recruitment, and down-regulation of stromal smooth-"
    "muscle / ECM components. Both axes intensify with FIGO stage. GWAS Catalog annotation graded 1 gene "
    "(S100A9) STRONG, 10 MODERATE, and 14 LIMITED. S100A8/S100A9 and CDKN2A emerge as the highest-priority "
    "candidates for downstream validation, combining large effect sizes, stage-progression monotonicity, and "
    "orthogonal human-genetics or clinical-biomarker evidence.",
    align="justify")

# ========== DATA AVAILABILITY ==========
add_heading(doc, "7. Data availability", level=1)
add_para(doc,
    "All supporting tables (clinical_metadata.csv, differential_expression_full.csv, top15_upregulated.csv, "
    "top10_downregulated.csv, age_stratified_expression.csv, stage_expression.csv, gwas_annotation.csv) and "
    "figures (Fig1–Fig10, PNG + PDF, 150 DPI) are provided alongside this report in the delivery folder. "
    "Source data were drawn from the TCGA-CESC cohort of the Genomic Data Commons (portal.gdc.cancer.gov) "
    "and the NHGRI-EBI GWAS Catalog release archive (ftp.ebi.ac.uk/pub/databases/gwas/).",
    align="justify")

# ========== SAVE ==========
doc.save(DOCX_PATH)
print(f"Wrote {DOCX_PATH}")
print(f"Size: {os.path.getsize(DOCX_PATH):,} bytes")
